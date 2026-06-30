"""Internal Knowledge Base — document ingestion + retrieval (RAG) service.

The retrieval layer is deterministic and fully offline (TF-IDF + cosine over the
user's own chunks), so search and the *retrieved* context never depend on an
external API. Q&A and summarization opportunistically use the configured LLM to
synthesise an answer, but always degrade gracefully to an extractive/template
result when ``llm.is_configured()`` is False or the provider raises
``LLMUnavailable``.

Heavy parsing libraries (pypdf / python-docx / openpyxl) are imported lazily so
this module imports cleanly even when those optional dependencies are absent.
"""

from __future__ import annotations

import re
from pathlib import Path

from sqlalchemy import select
from sqlalchemy.orm import Session

from mirrorquant_demo import llm
from mirrorquant_demo.kb_models import Document, DocumentChunk

DATA_DIR = Path(__file__).resolve().parent / "data"

_SUPPORTED_KINDS = {"pdf", "docx", "xlsx", "txt", "md", "text"}
_PLAINTEXT_KINDS = {"txt", "md", "text"}


# --- text extraction --------------------------------------------------------
def extract_text(path_or_bytes, kind: str) -> str:
    """Extract plain text from a file path / bytes by ``kind``.

    Heavy libraries are imported lazily; if a required parser is not installed a
    clear ``ValueError('install <lib> to ingest <kind>')`` is raised so the
    caller can surface a 400 rather than crashing.
    """
    kind = (kind or "").strip().lower()
    if kind not in _SUPPORTED_KINDS:
        raise ValueError(f"unsupported document kind: {kind!r}")

    if kind in _PLAINTEXT_KINDS:
        return _read_plaintext(path_or_bytes)

    if kind == "pdf":
        return _extract_pdf(path_or_bytes)
    if kind == "docx":
        return _extract_docx(path_or_bytes)
    if kind == "xlsx":
        return _extract_xlsx(path_or_bytes)

    # Defensive: unreachable given the guard above.
    raise ValueError(f"unsupported document kind: {kind!r}")


def _read_plaintext(path_or_bytes) -> str:
    if isinstance(path_or_bytes, (bytes, bytearray)):
        return bytes(path_or_bytes).decode("utf-8", errors="replace")
    if isinstance(path_or_bytes, str) and "\n" not in path_or_bytes and Path(path_or_bytes).exists():
        return Path(path_or_bytes).read_text(encoding="utf-8", errors="replace")
    if isinstance(path_or_bytes, Path):
        return path_or_bytes.read_text(encoding="utf-8", errors="replace")
    # Treat a raw string as the text itself.
    return str(path_or_bytes)


def _as_binary_stream(path_or_bytes):
    """Return an object the parser libraries accept (path str or BytesIO)."""
    if isinstance(path_or_bytes, (bytes, bytearray)):
        import io

        return io.BytesIO(bytes(path_or_bytes))
    return path_or_bytes


def _extract_pdf(path_or_bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise ValueError("install pypdf to ingest pdf") from exc

    reader = PdfReader(_as_binary_stream(path_or_bytes))
    parts = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(parts).strip()


def _extract_docx(path_or_bytes) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise ValueError("install python-docx to ingest docx") from exc

    doc = DocxDocument(_as_binary_stream(path_or_bytes))
    parts = [para.text for para in doc.paragraphs if para.text]
    return "\n".join(parts).strip()


def _extract_xlsx(path_or_bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise ValueError("install openpyxl to ingest xlsx") from exc

    workbook = load_workbook(_as_binary_stream(path_or_bytes), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"# {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if cell is None else str(cell) for cell in row]
            if any(cell.strip() for cell in cells):
                lines.append("\t".join(cells))
    try:
        workbook.close()
    except Exception:  # pragma: no cover - best-effort cleanup
        pass
    return "\n".join(lines).strip()


# --- chunking ---------------------------------------------------------------
def chunk_text(text: str, max_chars: int = 1200, overlap: int = 150) -> list[str]:
    """Split ``text`` into overlapping character windows.

    Tries to break on whitespace near the window boundary to avoid splitting
    words. Deterministic and dependency-free.
    """
    text = (text or "").strip()
    if not text:
        return []
    if max_chars <= 0:
        return [text]
    overlap = max(0, min(overlap, max_chars - 1))

    chunks: list[str] = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + max_chars, length)
        if end < length:
            # Prefer to break on the last whitespace within the window.
            window = text[start:end]
            split = window.rfind(" ")
            newline = window.rfind("\n")
            split = max(split, newline)
            if split > max_chars * 0.5:
                end = start + split
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= length:
            break
        start = max(end - overlap, start + 1)
    return chunks


# --- ingestion + CRUD -------------------------------------------------------
def ingest_document(
    session: Session,
    user_id: int,
    title: str,
    text: str,
    kind: str = "text",
    source_name: str | None = None,
) -> Document:
    """Persist a document and its chunks; returns the stored ``Document``."""
    text = text or ""
    document = Document(
        user_id=user_id,
        title=title,
        source_name=source_name,
        kind=(kind or "text"),
        char_count=len(text),
    )
    session.add(document)
    session.flush()  # assign document.id

    for index, chunk in enumerate(chunk_text(text)):
        session.add(
            DocumentChunk(document_id=document.id, chunk_index=index, text=chunk)
        )

    session.commit()
    session.refresh(document)
    return document


def list_documents(session: Session, user_id: int) -> list[Document]:
    stmt = (
        select(Document)
        .where(Document.user_id == user_id)
        .order_by(Document.created_at.desc())
    )
    return list(session.scalars(stmt).all())


def get_document(session: Session, user_id: int, doc_id: int) -> Document:
    document = session.get(Document, doc_id)
    if document is None or document.user_id != user_id:
        raise LookupError(f"document {doc_id} not found")
    return document


# --- retrieval --------------------------------------------------------------
def _user_chunks(session: Session, user_id: int) -> list[DocumentChunk]:
    stmt = (
        select(DocumentChunk)
        .join(Document, DocumentChunk.document_id == Document.id)
        .where(Document.user_id == user_id)
        .order_by(DocumentChunk.document_id, DocumentChunk.chunk_index)
    )
    return list(session.scalars(stmt).all())


def search_chunks(
    session: Session, user_id: int, query: str, top_k: int = 5
) -> list[dict]:
    """TF-IDF + cosine retrieval over the user's own chunks. Deterministic, offline.

    Returns ``[]`` when the user has no chunks. sklearn is imported lazily.
    """
    query = (query or "").strip()
    chunks = _user_chunks(session, user_id)
    if not chunks or not query:
        return []

    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise ValueError("install scikit-learn to search the knowledge base") from exc

    corpus = [chunk.text for chunk in chunks]
    vectorizer = TfidfVectorizer(stop_words="english")
    try:
        matrix = vectorizer.fit_transform(corpus)
        query_vec = vectorizer.transform([query])
    except ValueError:
        # Empty vocabulary (e.g. only stop-words / punctuation).
        return []

    scores = cosine_similarity(query_vec, matrix)[0]

    # Build doc-title lookup so we avoid lazy-loading per chunk.
    doc_ids = {chunk.document_id for chunk in chunks}
    titles = {
        doc_id: title
        for doc_id, title in session.execute(
            select(Document.id, Document.title).where(Document.id.in_(doc_ids))
        ).all()
    }

    ranked = sorted(
        range(len(chunks)), key=lambda i: float(scores[i]), reverse=True
    )
    top_k = max(0, int(top_k or 0))
    results: list[dict] = []
    for i in ranked[:top_k]:
        score = float(scores[i])
        if score <= 0.0:
            continue
        chunk = chunks[i]
        results.append(
            {
                "document_id": chunk.document_id,
                "document_title": titles.get(chunk.document_id, ""),
                "chunk_index": chunk.chunk_index,
                "text": chunk.text,
                "score": round(score, 6),
            }
        )
    return results


# --- Q&A + summarization ----------------------------------------------------
_ANSWER_SYSTEM = (
    "You are a research assistant for an internal financial knowledge base. "
    "Answer the user's question using ONLY the provided context excerpts. "
    "Cite the document titles you used inline. If the context does not contain "
    "the answer, say so plainly. Do not give investment advice or buy/sell "
    "recommendations."
)


def _format_context(retrieved: list[dict]) -> str:
    blocks = []
    for item in retrieved:
        label = f"[{item['document_title']} #chunk{item['chunk_index']}]"
        blocks.append(f"{label}\n{item['text']}")
    return "\n\n".join(blocks)


async def answer_question(
    session: Session, user_id: int, query: str, top_k: int = 5
) -> dict:
    """Retrieve top chunks then synthesise an answer (LLM if configured, else extractive)."""
    retrieved = search_chunks(session, user_id, query, top_k=top_k)
    sources = [
        {"document_title": item["document_title"], "chunk_index": item["chunk_index"]}
        for item in retrieved
    ]

    if not retrieved:
        return {
            "query": query,
            "answer": "No relevant documents were found in the knowledge base.",
            "llm_used": False,
            "sources": sources,
            "retrieved": retrieved,
        }

    context = _format_context(retrieved)

    if llm.is_configured():
        try:
            user_prompt = (
                f"Question:\n{query}\n\n"
                f"Context excerpts:\n{context}\n\n"
                "Answer the question and cite the document titles you used."
            )
            answer = await llm.complete(
                _ANSWER_SYSTEM, user_prompt, temperature=0.3, max_tokens=1500
            )
            answer = (answer or "").strip()
            if answer:
                return {
                    "query": query,
                    "answer": answer,
                    "llm_used": True,
                    "sources": sources,
                    "retrieved": retrieved,
                }
        except llm.LLMUnavailable:
            pass  # fall through to deterministic answer

    # Deterministic fallback: concatenate the retrieved chunks.
    titles = ", ".join(sorted({item["document_title"] for item in retrieved}))
    note = (
        "No LLM is configured; returning the most relevant excerpts verbatim "
        f"from: {titles}."
    )
    answer = note + "\n\n" + "\n\n---\n\n".join(item["text"] for item in retrieved)
    return {
        "query": query,
        "answer": answer,
        "llm_used": False,
        "sources": sources,
        "retrieved": retrieved,
    }


_SUMMARY_SYSTEM = (
    "You are a research analyst. Produce a concise, faithful summary of the "
    "provided document. Capture the key points only; do not invent facts and do "
    "not give investment advice or buy/sell recommendations."
)

_SENTENCE_SPLIT = re.compile(r"(?<=[.!?。！？])\s+")


def _extractive_summary(text: str, max_sentences: int = 6) -> str:
    text = (text or "").strip()
    if not text:
        return ""
    sentences = [s.strip() for s in _SENTENCE_SPLIT.split(text) if s.strip()]
    if not sentences:
        return text[:600]
    return " ".join(sentences[:max_sentences])


async def summarize_document(session: Session, user_id: int, doc_id: int) -> dict:
    """Summarize a document (LLM if configured, else first ~6 sentences)."""
    document = get_document(session, user_id, doc_id)
    full_text = "\n\n".join(chunk.text for chunk in document.chunks).strip()

    if full_text and llm.is_configured():
        try:
            summary = await llm.complete(
                _SUMMARY_SYSTEM,
                f"Document title: {document.title}\n\n{full_text}",
                temperature=0.3,
                max_tokens=1500,
            )
            summary = (summary or "").strip()
            if summary:
                return {
                    "document_id": document.id,
                    "title": document.title,
                    "summary": summary,
                    "llm_used": True,
                }
        except llm.LLMUnavailable:
            pass  # fall through to extractive summary

    return {
        "document_id": document.id,
        "title": document.title,
        "summary": _extractive_summary(full_text),
        "llm_used": False,
    }
